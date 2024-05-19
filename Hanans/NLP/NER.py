import fitz  # PyMuPDF
from docx import Document
from docx.shared import RGBColor
import spacy
from spacy import displacy
from langdetect import detect
from googletrans import Translator
import pandas as pd
from termcolor import colored

# Load SpaCy models
nlp_en = spacy.load("en_core_web_sm")
nlp_ur = spacy.load("xx_ent_wiki_sm")  # Multilingual model, as there's no dedicated Urdu model
translator = Translator()

def detect_language(text):
    return detect(text)

def translate_to_english(text):
    return translator.translate(text, src='ur', dest='en').text

def translate_to_urdu(text):
    return translator.translate(text, src='en', dest='ur').text

def apply_ner(text, language):
    if language == 'en':
        doc = nlp_en(text)
    else:
        doc = nlp_ur(text)
    entities = [{'text': ent.text, 'label': ent.label_} for ent in doc.ents]
    return entities

def highlight_entities(text, entities):
    highlighted_text = text
    for entity in entities:
        highlighted_text = highlighted_text.replace(entity['text'], colored(entity['text'], 'blue', attrs=['bold']))
    return highlighted_text

def entities_to_dataframe(entities):
    return pd.DataFrame(entities)

def process_text(text):
    language = detect_language(text)
    original_text = text

    if language == 'ur':
        text = translate_to_english(text)

    entities = apply_ner(text, 'en')

    if language == 'ur':
        translated_entities = []
        for entity in entities:
            translated_text = translate_to_urdu(entity['text'])
            translated_entities.append({'text': translated_text, 'label': entity['label']})
        entities = translated_entities

    highlighted_text = highlight_entities(original_text, entities)
    df = entities_to_dataframe(entities)

    return language, df, highlighted_text

def display_ner_visualization(text, language):
    if language == 'en':
        doc = nlp_en(text)
    else:
        doc = nlp_ur(text)
    displacy.render(doc, style='ent', jupyter=False)

def clean_text(text):
    # Remove non-printable and non-XML compatible characters
    return ''.join(c for c in text if c.isprintable() and c not in '\x00-\x1f\x7f-\x9f')

def extract_ne(pdf_path, output_path):
    pdf_document = fitz.open(pdf_path)
    doc = Document()

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text = page.get_text()

        language, df, highlighted_text = process_text(text)

        paragraphs = highlighted_text.split('\n')

        for paragraph_text in paragraphs:
            paragraph = doc.add_paragraph()
            words = paragraph_text.split()

            for word in words:
                cleaned_word = clean_text(word + ' ')
                run = paragraph.add_run(cleaned_word)
                run.font.color.rgb = RGBColor(0, 0, 255)

    doc.save(output_path)
    print("Named entities extracted and saved to", output_path)
