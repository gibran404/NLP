


import requests
import fitz  # PyMuPDF
from docx import Document
from docx.shared import RGBColor
import spacy
from spacy import displacy
from langdetect import detect

# Load SpaCy models
nlp_en = spacy.load("en_core_web_sm")
nlp_ur = spacy.load("xx_ent_wiki_sm")  # Multilingual model, as there's no dedicated Urdu model

LIBRETRANSLATE_URL = "https://libretranslate.com/translate"

def translate_text(text, source_language, target_language):
    response = requests.post(LIBRETRANSLATE_URL, data={
        "q": text,
        "source": source_language,
        "target": target_language,
        "format": "text"
    })
    result = response.json()
    return result["translatedText"]

def detect_language(text):
    return detect(text)

def translate_to_english(text):
    return translate_text(text, 'ur', 'en')

def translate_to_urdu(text):
    return translate_text(text, 'en', 'ur')

def apply_ner(text, language):
    if language == 'en':
        doc = nlp_en(text)
    else:
        doc = nlp_ur(text)
    entities = [{'text': ent.text, 'label': ent.label_} for ent in doc.ents]
    return entities

def get_entities(text):
    language = detect_language(text)
    original_text = text

    if language == 'ur':
        text = translate_to_english(text)

    entities = apply_ner(text, 'en')

    if language == 'ur':
        translated_entities = []
        for entity in entities:
            translated_text = translate_to_urdu(entity['text'])
            translated_entities.append({'text': translated_text, 'label': entity['label']})
        entities = translated_entities

    return entities

def get_color(label):
    colors = {
        'PERSON': RGBColor(255, 0, 0),
        'NORP': RGBColor(0, 255, 0),
        'FAC': RGBColor(0, 0, 255),
        'ORG': RGBColor(255, 255, 0),
        'GPE': RGBColor(255, 0, 255),
        'LOC': RGBColor(0, 255, 255),
        'PRODUCT': RGBColor(255, 192, 203),
        'EVENT': RGBColor(139, 69, 19),
        'WORK_OF_ART': RGBColor(169, 169, 169),
        'LAW': RGBColor(0, 0, 0),
        'LANGUAGE': RGBColor(255, 255, 255),
        'DATE': RGBColor(0, 255, 255),
        'TIME': RGBColor(255, 0, 255),
        'PERCENT': RGBColor(211, 211, 211),
        'MONEY': RGBColor(105, 105, 105),
        'QUANTITY': RGBColor(173, 216, 230),
        'ORDINAL': RGBColor(144, 238, 144),
        'CARDINAL': RGBColor(255, 255, 224)
    }
    return colors.get(label, RGBColor(0, 0, 0))

def extract_ne(pdf_path, output_path):
    pdf_document = fitz.open(pdf_path)
    doc = Document()

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text = page.get_text()

        entities_list = get_entities(text)
        current_word = ''
        
        paragraph = doc.add_paragraph()
        
        for char in text:
            if char.isalpha():
                current_word += char
            else:
                if current_word:
                    matched_entity = next((entity for entity in entities_list if entity['text'] == current_word), None)
                    if matched_entity:
                        run = paragraph.add_run(current_word)
                        run.font.color.rgb = get_color(matched_entity['label'])
                    else:
                        paragraph.add_run(current_word)
                    current_word = ''
                if char.isspace():
                    paragraph.add_run(char)
                else:
                    run = paragraph.add_run(char)
                    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(output_path)
    print("Named entities extracted and saved to", output_path)

extract_ne('urdu.pdf', 'urduoutput.docx')
