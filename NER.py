
import fitz  # PyMuPDF
from docx import Document
from docx.shared import RGBColor
import spacy
from spacy import displacy
from langdetect import detect
from googletrans import Translator

# Load SpaCy models
nlp_en = spacy.load("en_core_web_sm")
nlp_ur = spacy.load("xx_ent_wiki_sm")  # Multilingual model, as there's no dedicated Urdu model
translator = Translator()

def detect_language(text):
    return detect(text)

def translate_to_english(text):
    return translator.translate(text, src='ur', dest='en').text

def translate_to_urdu(text):
    return translator.translate(text, src='en', dest='ur').text

def apply_ner(text, language):
    if language == 'en':
        doc = nlp_en(text)
    else:
        doc = nlp_ur(text)
    entities = [{'text': ent.text, 'label': ent.label_} for ent in doc.ents]
    return entities

def get_entities(text):
    language = detect_language(text)
    original_text = text

    if language == 'ur':
        text = translate_to_english(text)

    entities = apply_ner(text, 'en')

    if language == 'ur':
        translated_entities = []
        for entity in entities:
            translated_text = translate_to_urdu(entity['text'])
            translated_entities.append({'text': translated_text, 'label': entity['label']})
        entities = translated_entities

    return entities

def get_color(label):
    colors = {
        'PERSON': RGBColor(255, 0, 0), # red
        'NORP': RGBColor(0, 255, 0), # green
        'FAC': RGBColor(0, 0, 255), # blue
        'ORG': RGBColor(171, 171, 0), # dark yellow
        'GPE': RGBColor(255, 0, 255), # purple
        'LOC': RGBColor(0, 206, 206), # cyan
        'PRODUCT': RGBColor(195, 150, 158), # pink
        'EVENT': RGBColor(139, 69, 19),     # brown
        'WORK_OF_ART': RGBColor(169, 169, 169), # dark gray
        'LAW': RGBColor(153, 0, 0), # dark red
        'LANGUAGE': RGBColor(0, 102, 102),
        'DATE': RGBColor(0, 102, 0), # dark green
        'TIME': RGBColor(255, 0, 255), # magenta
        'PERCENT': RGBColor(211, 211, 211), # light gray
        'MONEY': RGBColor(105, 105, 105), # dark gray
        'QUANTITY': RGBColor(173, 216, 230), # light blue
        'ORDINAL': RGBColor(144, 238, 144), # light green
        'CARDINAL': RGBColor(0, 0, 102) # light yellow
    }
    return colors.get(label, RGBColor(0, 0, 0))

def extract_ne(pdf_path, output_path):
    pdf_document = fitz.open(pdf_path)
    doc = Document()
    
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text = page.get_text()

        entities_list = get_entities(text)
        entities_dict = {entity['text']: entity['label'] for entity in entities_list}
        
        paragraph = doc.add_paragraph()
        
        pos = 0
        while pos < len(text):
            found = False
            for entity_text, entity_label in entities_dict.items():
                if text[pos:pos + len(entity_text)] == entity_text:
                    # Add the entity with color and bold formatting
                    run = paragraph.add_run(entity_text)
                    run.font.color.rgb = get_color(entity_label)
                    run.bold = True

                    run = paragraph.add_run(f" ({entity_label})")
                    run.font.color.rgb = get_color(entity_label)
                    run.bold = True

                    pos += len(entity_text)
                    found = True
                    break
            
            if not found:
                # Add the current character and move to the next one
                run = paragraph.add_run(text[pos])
                run.font.color.rgb = RGBColor(0, 0, 0)
                pos += 1

    doc.save(output_path)
    print("Named entities extracted and saved to", output_path)


